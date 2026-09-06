import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_element(name):
    return OxmlElement(name)

def add_bottom_border(paragraph, color_hex="1E3A8A", sz="10"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = create_element('w:pBdr')
    bottom = create_element('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), sz) # 10 = 1.25 pt
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def build_resume_v2():
    doc = docx.Document()
    
    # Page Margins (0.5 in top/bottom, 0.6 in left/right)
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
    
    # Palette definition
    COLOR_PRIMARY = RGBColor(15, 23, 42)      # Deep Slate Navy #0F172A
    COLOR_ACCENT = RGBColor(30, 58, 138)      # Deep Royal Blue #1E3A8A
    COLOR_BODY = RGBColor(51, 65, 85)         # Slate Charcoal #334155
    COLOR_MUTED = RGBColor(100, 116, 139)     # Cool Slate Grey #64748B
    
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(9.5)
    normal_style.font.color.rgb = COLOR_BODY
    
    # --- HEADER SECTION ---
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(2)
    run_name = p_name.add_run("NIDHISH V")
    run_name.font.size = Pt(20)
    run_name.font.bold = True
    run_name.font.color.rgb = COLOR_PRIMARY
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("CYBERSECURITY DEFENSE EXPERT | THREAT ANALYSIS & AI SECURITY EVALUATOR")
    run_title.font.size = Pt(10)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_ACCENT
    
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_before = Pt(0)
    p_contact.paragraph_format.space_after = Pt(10)
    run_contact = p_contact.add_run("Chennai, Tamil Nadu, India  •  +91 8056209809  •  nidhishv31@gmail.com  •  github.com/nidhishv31-cpu")
    run_contact.font.size = Pt(9)
    run_contact.font.color.rgb = COLOR_MUTED
    
    def add_section_header(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(9)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(title.upper())
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        add_bottom_border(p, color_hex="1E3A8A", sz="10")
        return p

    def add_bullet(p, bold_prefix, text):
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2.5)
        p.paragraph_format.line_spacing = 1.14
        
        run_b = p.add_run("▪  ")
        run_b.font.size = Pt(8.5)
        run_b.font.bold = True
        run_b.font.color.rgb = COLOR_ACCENT
        
        if bold_prefix:
            run_p = p.add_run(bold_prefix)
            run_p.font.bold = True
            run_p.font.size = Pt(9.5)
            run_p.font.color.rgb = COLOR_PRIMARY
        
        run_t = p.add_run(text)
        run_t.font.size = Pt(9.5)
        run_t.font.color.rgb = COLOR_BODY

    # --- PROFESSIONAL SUMMARY ---
    add_section_header("Professional Summary")
    p_sum = doc.add_paragraph()
    p_sum.paragraph_format.space_before = Pt(3)
    p_sum.paragraph_format.space_after = Pt(6)
    p_sum.paragraph_format.line_spacing = 1.15
    run_sum = p_sum.add_run(
        "Methodical, defense-focused Cybersecurity Specialist with deep technical grounding in network forensics, "
        "threat detection engineering, protocol inspection, and defensive rule validation. Possesses hands-on experience "
        "architecting enterprise-grade security platforms, writing and evaluating production detection logic (Suricata IDS, "
        "Linux iptables, Windows netsh firewall rules), dissecting packets down to raw bytes (Wireshark/Tshark), and auditing defensive "
        "Python/Bash automation scripts. Adept at applying the MITRE ATT&CK framework, Cyber Kill Chain, and NIST/CIS baselines "
        "to evaluate attack vectors and validate defensive measures. Equipped with the analytical rigor and domain instincts needed "
        "to critically evaluate, stress-test, and benchmark AI-generated cybersecurity content—identifying subtle security "
        "misconfigurations, outdated practices, syntax errors, and dangerous hallucinations to ensure production-safe, actionable AI guidance."
    )
    run_sum.font.size = Pt(9.5)
    run_sum.font.color.rgb = COLOR_BODY

    # --- CORE TECHNICAL COMPETENCIES ---
    add_section_header("Core Technical Competencies")
    skills_data = [
        ("Threat Frameworks & Defensive Strategy: ", "MITRE ATT&CK Framework (v14 TTP mapping), Cyber Kill Chain, NIST Cybersecurity Framework (CSF), CIS Controls, OWASP Top 10 Web Vulnerabilities, Incident Response (IR) Runbooks, Threat Modeling."),
        ("Detection Engineering & Rule Auditing: ", "Suricata IDS Rule Syntax & Production Viability, Perimeter Firewall Policies (Linux iptables, Windows netsh advfirewall), YARA Heuristic Artifact Carving, OASIS STIX 2.1 Threat Intel Bundles, SIEM/SOC Telemetry Correlation."),
        ("AI Security Content Evaluation: ", "Critical Stress-Testing of AI-Generated Defensive Guidance, Hallucination & Outdated Practice Identification, Defensive Code Hardening (Python, Bash, Sockets), Misconfiguration Detection, Structured Technical Advisory Writing."),
        ("Network Forensics & Deep Packet Inspection: ", "Deep Packet Inspection (DPI), Wireshark & Tshark Packet Dissection, RFC 1071 Checksum Verification, DNS Tunneling & DGA Statistical Detection (Shannon Entropy), In-Flight TLS Decryption (SSLKEYLOGFILE), TCP/IP Socket Auditing."),
        ("Security Tools & Platforms: ", "Wireshark, Tshark, Nmap / Zenmap, Burp Suite, Metasploit Framework, Kali Linux Security Suite, Git / GitHub."),
        ("Programming & Automation: ", "Python 3 (Scapy, Cryptography, Requests, Sockets), Bash Shell Scripting, Node.js / JavaScript basics, C, C++."),
        ("Operating Systems & Administration: ", "Linux (Kali Linux, Ubuntu, Debian CLI Administration, Syscall Auditing), Microsoft Windows (10/11, PowerShell, Command Line), MS Office Suite.")
    ]
    for cat, items in skills_data:
        p_s = doc.add_paragraph()
        add_bullet(p_s, cat, items)

    # --- DEFENSIVE SECURITY & RESEARCH PROJECTS ---
    add_section_header("Defensive Security & Engineering Projects")
    
    # Project 1: SENTINEL ENTERPRISE / VULNSCAN
    p_proj1 = doc.add_paragraph()
    p_proj1.paragraph_format.space_before = Pt(4)
    p_proj1.paragraph_format.space_after = Pt(1)
    p_proj1.paragraph_format.keep_with_next = True
    r_p1_title = p_proj1.add_run("SENTINEL ENTERPRISE / VULNSCAN — Network Forensics, Threat Hunting & AI SOC Platform")
    r_p1_title.font.bold = True
    r_p1_title.font.size = Pt(10)
    r_p1_title.font.color.rgb = COLOR_ACCENT
    
    p_p1_tech = doc.add_paragraph()
    p_p1_tech.paragraph_format.space_before = Pt(0)
    p_p1_tech.paragraph_format.space_after = Pt(3)
    p_p1_tech.paragraph_format.keep_with_next = True
    r_p1_tech = p_p1_tech.add_run("Technologies: Python (FastAPI), React 19, Tshark/Wireshark, Scapy, YARA, OASIS STIX 2.1, MITRE ATT&CK v14")
    r_p1_tech.font.size = Pt(8.5)
    r_p1_tech.font.italic = True
    r_p1_tech.font.color.rgb = COLOR_MUTED

    sentinel_bullets = [
        "Defensive Rule Generation & Production Validation: Built an AI SOC Incident Brief & Rule Engine synthesizing and auditing production-ready perimeter defense rules—including Linux iptables, Windows netsh advfirewall, and Suricata IDS alerts directly mapped to MITRE ATT&CK TTPs (e.g., T1071.004 DNS Tunneling).",
        "Threat Detection & Statistical Heuristics: Engineered multi-layered detection models utilizing Shannon Entropy (H(X) >= 3.8) to mathematically differentiate cleartext traffic from high-entropy obfuscated/encrypted C2 beacons, combined with YARA heuristic artifact carving for disguised PE binaries.",
        "Protocol RFC Inspection & Packet Dissection: Implemented Layer 2-4 frame dissection and replay injection verified against RFC 1071 16-bit Internet Checksums, ensuring synthesized packets and defensive responses conform strictly to standards without causing network stack instability.",
        "Self-Auditing Pipeline via Wireshark / Tshark: Integrated a live packet-capture engine to self-audit outbound scanner and detection traffic, verifying rate-limit enforcement, payload structural integrity, and zero plaintext credential or token leakage.",
        "OASIS STIX 2.1 Threat Intel Modeling: Formatted and validated high-confidence Indicators of Compromise (IoCs) into RFC-compliant STIX 2.1 JSON bundles, verifying interoperability for enterprise SIEM ingestion (Splunk, Microsoft Sentinel, Elastic)."
    ]
    for b in sentinel_bullets:
        p_b = doc.add_paragraph()
        add_bullet(p_b, "", b)

    # Project 2: STEGTOOL
    p_proj2 = doc.add_paragraph()
    p_proj2.paragraph_format.space_before = Pt(5)
    p_proj2.paragraph_format.space_after = Pt(1)
    p_proj2.paragraph_format.keep_with_next = True
    r_p2_title = p_proj2.add_run("STEGTOOL — Digital Forensics, Cryptographic Verification & Steganalysis Suite")
    r_p2_title.font.bold = True
    r_p2_title.font.size = Pt(10)
    r_p2_title.font.color.rgb = COLOR_ACCENT
    
    p_p2_tech = doc.add_paragraph()
    p_p2_tech.paragraph_format.space_before = Pt(0)
    p_p2_tech.paragraph_format.space_after = Pt(3)
    p_p2_tech.paragraph_format.keep_with_next = True
    r_p2_tech = p_p2_tech.add_run("Technologies: Python, AES-256-GCM Cryptography, SciPy, OpenCV/PIL, ReportLab")
    r_p2_tech.font.size = Pt(8.5)
    r_p2_tech.font.italic = True
    r_p2_tech.font.color.rgb = COLOR_MUTED

    stegtool_bullets = [
        "Cryptographic Rigor & Data Defense: Developed an enterprise-grade digital steganography and forensic analysis suite providing secure Least Significant Bit (LSB) payload embedding across image (PNG/BMP) and audio (WAV) carriers, reinforced with authenticated AES-256-GCM encryption and mandatory 2FA keyfile verification.",
        "Mathematical Steganalysis & Anomaly Detection: Engineered a statistical Chi-Square (Chi^2) steganalysis engine to evaluate carrier sample distributions, spotting hidden payload artifacts and detecting covert channel exfiltrations that evade signature-based detection.",
        "Structured Forensic Technical Reporting: Implemented automated generation of comprehensive forensic audit reports providing clear, structured technical breakdowns of sample entropy, payload integrity, and carrier health for incident investigations."
    ]
    for b in stegtool_bullets:
        p_b = doc.add_paragraph()
        add_bullet(p_b, "", b)

    # --- AI EVALUATION & STRESS-TESTING STRENGTHS ---
    add_section_header("AI Training & Security Evaluation Fit")
    fit_bullets = [
        ("Detection Rule & Configuration Verification: ", "Hands-on experience authoring and testing Suricata, iptables, and netsh configurations ensures rapid identification of syntax bugs, conflicting rules, or ineffective perimeter policies in AI responses."),
        ("Anti-Hallucination & RFC-Grade Technical Precision: ", "Grounded in network RFC specifications (RFC 1071 checksums, TCP handshakes, TLS negotiation), enabling instant detection of technically invalid parameters or dangerous hallucinations."),
        ("Threat Taxonomy Alignment: ", "Proven ability to map attack scenarios and mitigations against the MITRE ATT&CK matrix and Cyber Kill Chain, ensuring AI security recommendations represent industry standard defense."),
        ("Clear Technical Advisory Communication: ", "Track record of producing comprehensive security documentation and structured technical evaluations explaining exactly why a recommended configuration is insecure and how to remediate it.")
    ]
    for pref, body in fit_bullets:
        p_f = doc.add_paragraph()
        add_bullet(p_f, pref, body)

    # --- EDUCATION ---
    add_section_header("Education")
    
    # SASTRA
    p_edu1 = doc.add_paragraph()
    p_edu1.paragraph_format.space_before = Pt(3)
    p_edu1.paragraph_format.space_after = Pt(1)
    p_edu1.paragraph_format.keep_with_next = True
    r_deg = p_edu1.add_run("Bachelor of Technology (B.Tech) – Electronics and Communication Engineering")
    r_deg.font.bold = True
    r_deg.font.size = Pt(10)
    r_deg.font.color.rgb = COLOR_PRIMARY
    
    p_inst1 = doc.add_paragraph()
    p_inst1.paragraph_format.space_before = Pt(0)
    p_inst1.paragraph_format.space_after = Pt(2)
    r_inst1 = p_inst1.add_run("SASTRA Deemed University, Thanjavur, Tamil Nadu  |  Aug 2024 – Present")
    r_inst1.font.italic = True
    r_inst1.font.size = Pt(9)
    r_inst1.font.color.rgb = COLOR_MUTED

    p_e1_det = doc.add_paragraph()
    add_bullet(p_e1_det, "Academic Standing: ", "CGPA: 6.5 / 10.0  |  Specialized Elective: Cyber Security")
    p_e1_crs = doc.add_paragraph()
    add_bullet(p_e1_crs, "Relevant Coursework: ", "Computer Communication Networks, Cyber Security Fundamentals, Operating Systems & Security, Digital Forensics, Cryptography, Data Structures.")

    # 12th Narayana Olympiad
    p_edu2 = doc.add_paragraph()
    p_edu2.paragraph_format.space_before = Pt(3)
    p_edu2.paragraph_format.space_after = Pt(1)
    p_edu2.paragraph_format.keep_with_next = True
    r_deg2 = p_edu2.add_run("Higher Secondary Certificate (Class XII)")
    r_deg2.font.bold = True
    r_deg2.font.size = Pt(10)
    r_deg2.font.color.rgb = COLOR_PRIMARY
    
    p_inst2 = doc.add_paragraph()
    p_inst2.paragraph_format.space_before = Pt(0)
    p_inst2.paragraph_format.space_after = Pt(2)
    r_inst2 = p_inst2.add_run("Narayana Olympiad School")
    r_inst2.font.italic = True
    r_inst2.font.size = Pt(9)
    r_inst2.font.color.rgb = COLOR_MUTED

    p_e2_det = doc.add_paragraph()
    add_bullet(p_e2_det, "Academic Score: ", "Aggregate: 67%  |  Focus: Physics, Chemistry, Mathematics & Computer Science")

    # 10th Narayana e-Techno
    p_edu3 = doc.add_paragraph()
    p_edu3.paragraph_format.space_before = Pt(3)
    p_edu3.paragraph_format.space_after = Pt(1)
    p_edu3.paragraph_format.keep_with_next = True
    r_deg3 = p_edu3.add_run("Secondary School Certificate (Class X)")
    r_deg3.font.bold = True
    r_deg3.font.size = Pt(10)
    r_deg3.font.color.rgb = COLOR_PRIMARY
    
    p_inst3 = doc.add_paragraph()
    p_inst3.paragraph_format.space_before = Pt(0)
    p_inst3.paragraph_format.space_after = Pt(2)
    r_inst3 = p_inst3.add_run("Narayana e-Techno School")
    r_inst3.font.italic = True
    r_inst3.font.size = Pt(9)
    r_inst3.font.color.rgb = COLOR_MUTED

    p_e3_det = doc.add_paragraph()
    add_bullet(p_e3_det, "Academic Score: ", "Aggregate: 76%  |  Science & Mathematics Foundation")

    # --- CERTIFICATIONS & PROFESSIONAL TRAINING ---
    add_section_header("Certifications & Technical Training")
    cert_data = [
        ("Google Cybersecurity Professional Certificate: ", "Security Operations Center (SOC) workflows, SIEM packet analysis, threat detection, Linux/SQL security automation, and incident response frameworks."),
        ("Microsoft Networking and Cloud Computing: ", "Enterprise network architecture, TCP/IP fundamentals, cloud security models, subnetting, firewalls, and cloud infrastructure management."),
        ("Mastering Operating Systems: ", "Process scheduling, memory virtualization, Unix/Linux system administration, privilege separation, and access control lists (ACLs)."),
        ("Cryptography & Network Security: ", "Hands-on implementation of symmetric/asymmetric encryption (AES, RSA), cryptographic hashing (SHA-256), public key infrastructure (PKI), and SSL/TLS security.")
    ]
    for c_title, c_desc in cert_data:
        p_c = doc.add_paragraph()
        add_bullet(p_c, c_title, c_desc)

    out_path = r"n:\Nidhish\project\Scanner\V2\Nidhish_V_Cybersecurity_Resume_V2.docx"
    doc.save(out_path)
    print(f"V2 Resume saved successfully to {out_path}")

if __name__ == "__main__":
    build_resume_v2()
