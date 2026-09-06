import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_element(name):
    return OxmlElement(name)

def add_bottom_border(paragraph, color_hex="1E3A8A", sz="12"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = create_element('w:pBdr')
    bottom = create_element('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), sz) # 12 = 1.5 pt
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def build_resume():
    doc = docx.Document()
    
    # Page Margins (0.55 in left/right, 0.5 in top/bottom for balanced ATS layout)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
    
    # Palette definition
    COLOR_PRIMARY = RGBColor(15, 23, 42)      # Deep Slate Navy #0F172A
    COLOR_ACCENT = RGBColor(30, 58, 138)      # Deep Blue #1E3A8A
    COLOR_BODY = RGBColor(51, 65, 85)         # Slate Charcoal #334155
    COLOR_MUTED = RGBColor(100, 116, 139)     # Cool Slate #64748B
    
    # Set default style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(9.5)
    normal_style.font.color.rgb = COLOR_BODY
    
    # --- HEADER SECTION ---
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(2)
    run_name = p_name.add_run("NIDHISH V")
    run_name.font.size = Pt(20)
    run_name.font.bold = True
    run_name.font.color.rgb = COLOR_PRIMARY
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("JUNIOR CYBERSECURITY ANALYST | SECURITY OPERATIONS (SOC) INTERN")
    run_title.font.size = Pt(10)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_ACCENT
    
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_before = Pt(0)
    p_contact.paragraph_format.space_after = Pt(10)
    run_contact = p_contact.add_run("Chennai, Tamil Nadu, India  •  +91 8056209809  •  nidhishv31@gmail.com  •  github.com/nidhishv31-cpu")
    run_contact.font.size = Pt(9)
    run_contact.font.color.rgb = COLOR_MUTED
    
    def add_section_header(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(9)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(title.upper())
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        add_bottom_border(p, color_hex="1E3A8A", sz="10")
        return p

    def add_bullet(p, bold_prefix, text):
        # Professional hanging indent: bullet sits at 0.10 in, text aligns at 0.25 in
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2.5)
        p.paragraph_format.line_spacing = 1.15
        
        run_b = p.add_run("▪  ")
        run_b.font.size = Pt(8.5)
        run_b.font.bold = True
        run_b.font.color.rgb = COLOR_ACCENT
        
        if bold_prefix:
            run_p = p.add_run(bold_prefix)
            run_p.font.bold = True
            run_p.font.size = Pt(9.5)
            run_p.font.color.rgb = COLOR_PRIMARY
        
        run_t = p.add_run(text)
        run_t.font.size = Pt(9.5)
        run_t.font.color.rgb = COLOR_BODY

    # --- CAREER OBJECTIVE ---
    add_section_header("Career Objective")
    p_obj = doc.add_paragraph()
    p_obj.paragraph_format.space_before = Pt(3)
    p_obj.paragraph_format.space_after = Pt(6)
    p_obj.paragraph_format.line_spacing = 1.15
    run_obj = p_obj.add_run(
        "Analytical and proactive Electronics & Communication Engineering undergraduate with a specialized elective in "
        "Cyber Security at SASTRA University, seeking a Junior Cybersecurity Analyst / Cybersecurity Intern role. "
        "Possesses hands-on technical proficiency in vulnerability assessment, protocol-level network analysis, TCP socket programming, "
        "and DevSecOps security tool development. Adept in utilizing Nmap, Wireshark, Burp Suite, Metasploit, and Python/Bash scripting "
        "to monitor security events, audit web applications against OWASP Top 10 vulnerabilities, evaluate network traffic, and support incident "
        "investigations. Dedicated to enforcing defensive security controls, mitigating system risks, and collaborating cross-functionally "
        "with engineering teams to safeguard critical digital assets."
    )
    run_obj.font.size = Pt(9.5)
    run_obj.font.color.rgb = COLOR_BODY

    # --- TECHNICAL SKILLS ---
    add_section_header("Technical Skills")
    
    skills_data = [
        ("Security & Vulnerability Assessment: ", "Vulnerability Scanning, Basic Penetration Testing, OWASP Top 10 Web Vulnerabilities, Incident Response Fundamentals, Security Auditing, Attack Surface Mapping, Risk Mitigation & Remediation."),
        ("Networking & Security Protocols: ", "TCP/IP Suite, DNS, HTTP/HTTPS, SSL/TLS Handshake & Cipher Inspection, Firewalls & Egress Filtering, Deep Packet Inspection (DPI), Network Socket Auditing."),
        ("Security Tools & Frameworks: ", "Wireshark, Tshark, Nmap / Zenmap, Burp Suite, Metasploit Framework, Kali Linux Security Suite."),
        ("Programming & Scripting: ", "Python (Sockets, Scapy, Requests, Scripting), C, C++, Bash Shell Scripting, JavaScript (Node.js)."),
        ("Operating Systems & Environments: ", "Linux (Kali Linux, Ubuntu, Debian CLI Administration), Microsoft Windows (10/11, Command Line, PowerShell basics), Git / GitHub."),
        ("Documentation & Professional Competencies: ", "Technical Vulnerability Reporting, Incident Documentation, Cross-Functional Teamwork, Analytical Problem-Solving, MS Office Suite.")
    ]
    for cat, items in skills_data:
        p_s = doc.add_paragraph()
        add_bullet(p_s, cat, items)

    # --- KEY CYBERSECURITY PROJECTS ---
    add_section_header("Key Cybersecurity Projects")
    
    # Project 1: VULNSCAN
    p_proj1 = doc.add_paragraph()
    p_proj1.paragraph_format.space_before = Pt(4)
    p_proj1.paragraph_format.space_after = Pt(1)
    p_proj1.paragraph_format.keep_with_next = True
    r_p1_title = p_proj1.add_run("VULNSCAN — Autonomous DevSecOps & Network Vulnerability Assessment Platform")
    r_p1_title.font.bold = True
    r_p1_title.font.size = Pt(10)
    r_p1_title.font.color.rgb = COLOR_ACCENT
    
    p_p1_tech = doc.add_paragraph()
    p_p1_tech.paragraph_format.space_before = Pt(0)
    p_p1_tech.paragraph_format.space_after = Pt(3)
    p_p1_tech.paragraph_format.keep_with_next = True
    r_p1_tech = p_p1_tech.add_run("Technologies: Python, Node.js (net.Socket), Wireshark/Tshark, Cheerio, REST APIs, Cryptography")
    r_p1_tech.font.size = Pt(8.5)
    r_p1_tech.font.italic = True
    r_p1_tech.font.color.rgb = COLOR_MUTED

    vulnscan_bullets = [
        "Engineered a full-stack DevSecOps vulnerability scanner and network auditing suite executing multi-vector security assessments across web applications and network endpoints.",
        "Automated low-level TCP socket port auditing (net.Socket), HTTP security header compliance analysis (HSTS, CSP, X-Frame-Options), and live SSL/TLS certificate chain & cipher suite validation.",
        "Integrated a Wireshark/Tshark packet-capture engine to self-audit outbound scanner traffic, enforcing rate-limit compliance, payload structural integrity, and zero plaintext credential/token leakage.",
        "Architected an automated web DOM integrity monitoring pipeline utilizing Node-cron and Cheerio to track real-time website defacement/tampering, log version diff history, and trigger instant SMTP/Webhook security alerts.",
        "Constructed an interactive analyst reporting dashboard that generates structured vulnerability assessment reports with severity ratings and actionable mitigation guidance aligned with OWASP best practices."
    ]
    for b in vulnscan_bullets:
        p_b = doc.add_paragraph()
        add_bullet(p_b, "", b)

    # Project 2: STEGTOOL
    p_proj2 = doc.add_paragraph()
    p_proj2.paragraph_format.space_before = Pt(5)
    p_proj2.paragraph_format.space_after = Pt(1)
    p_proj2.paragraph_format.keep_with_next = True
    r_p2_title = p_proj2.add_run("STEGTOOL — Enterprise Steganography & Digital Forensics Analysis Suite")
    r_p2_title.font.bold = True
    r_p2_title.font.size = Pt(10)
    r_p2_title.font.color.rgb = COLOR_ACCENT
    
    p_p2_tech = doc.add_paragraph()
    p_p2_tech.paragraph_format.space_before = Pt(0)
    p_p2_tech.paragraph_format.space_after = Pt(3)
    p_p2_tech.paragraph_format.keep_with_next = True
    r_p2_tech = p_p2_tech.add_run("Technologies: Python, AES-256-GCM Cryptography, SciPy, OpenCV/PIL, ReportLab (PDF Generation)")
    r_p2_tech.font.size = Pt(8.5)
    r_p2_tech.font.italic = True
    r_p2_tech.font.color.rgb = COLOR_MUTED

    stegtool_bullets = [
        "Developed an enterprise-grade digital steganography and forensic analysis suite providing secure Least Significant Bit (LSB) payload embedding across image (PNG, BMP) and audio (WAV) carriers.",
        "Fortified carrier data confidentiality and tamper-resistance utilizing authenticated AES-256-GCM encryption paired with mandatory Two-Factor Authentication (2FA) keyfile verification.",
        "Integrated a statistical Chi-Square (Chi^2) steganalysis engine to detect anomalous pixel distributions and hidden payload signatures, paired with automated PDF forensic report generation for threat investigations."
    ]
    for b in stegtool_bullets:
        p_b = doc.add_paragraph()
        add_bullet(p_b, "", b)

    # --- EDUCATION ---
    add_section_header("Education")
    
    # Sastra University
    p_edu1 = doc.add_paragraph()
    p_edu1.paragraph_format.space_before = Pt(3)
    p_edu1.paragraph_format.space_after = Pt(1)
    p_edu1.paragraph_format.keep_with_next = True
    r_deg = p_edu1.add_run("Bachelor of Technology (B.Tech) – Electronics and Communication Engineering")
    r_deg.font.bold = True
    r_deg.font.size = Pt(10)
    r_deg.font.color.rgb = COLOR_PRIMARY
    
    p_inst1 = doc.add_paragraph()
    p_inst1.paragraph_format.space_before = Pt(0)
    p_inst1.paragraph_format.space_after = Pt(2)
    r_inst1 = p_inst1.add_run("SASTRA Deemed University, Thanjavur, Tamil Nadu  |  Aug 2024 – Present")
    r_inst1.font.italic = True
    r_inst1.font.size = Pt(9)
    r_inst1.font.color.rgb = COLOR_MUTED

    p_e1_det = doc.add_paragraph()
    add_bullet(p_e1_det, "Academic Standing: ", "CGPA: 6.5 / 10.0  |  Specialized Elective: Cyber Security")
    p_e1_crs = doc.add_paragraph()
    add_bullet(p_e1_crs, "Core Coursework: ", "Computer Networks, Network Security, Operating Systems, Digital Forensics, Data Structures, Microprocessors & Sockets.")

    # 12th Narayana Olympiad
    p_edu2 = doc.add_paragraph()
    p_edu2.paragraph_format.space_before = Pt(3)
    p_edu2.paragraph_format.space_after = Pt(1)
    p_edu2.paragraph_format.keep_with_next = True
    r_deg2 = p_edu2.add_run("Higher Secondary Certificate (Class XII)")
    r_deg2.font.bold = True
    r_deg2.font.size = Pt(10)
    r_deg2.font.color.rgb = COLOR_PRIMARY
    
    p_inst2 = doc.add_paragraph()
    p_inst2.paragraph_format.space_before = Pt(0)
    p_inst2.paragraph_format.space_after = Pt(2)
    r_inst2 = p_inst2.add_run("Narayana Olympiad School")
    r_inst2.font.italic = True
    r_inst2.font.size = Pt(9)
    r_inst2.font.color.rgb = COLOR_MUTED

    p_e2_det = doc.add_paragraph()
    add_bullet(p_e2_det, "Academic Score: ", "Aggregate: 67%  |  Focus: Physics, Chemistry, Mathematics & Computer Science")

    # 10th Narayana e-Techno
    p_edu3 = doc.add_paragraph()
    p_edu3.paragraph_format.space_before = Pt(3)
    p_edu3.paragraph_format.space_after = Pt(1)
    p_edu3.paragraph_format.keep_with_next = True
    r_deg3 = p_edu3.add_run("Secondary School Certificate (Class X)")
    r_deg3.font.bold = True
    r_deg3.font.size = Pt(10)
    r_deg3.font.color.rgb = COLOR_PRIMARY
    
    p_inst3 = doc.add_paragraph()
    p_inst3.paragraph_format.space_before = Pt(0)
    p_inst3.paragraph_format.space_after = Pt(2)
    r_inst3 = p_inst3.add_run("Narayana e-Techno School")
    r_inst3.font.italic = True
    r_inst3.font.size = Pt(9)
    r_inst3.font.color.rgb = COLOR_MUTED

    p_e3_det = doc.add_paragraph()
    add_bullet(p_e3_det, "Academic Score: ", "Aggregate: 76%  |  Science & Mathematics Foundation")

    # --- CERTIFICATIONS & PROFESSIONAL COURSES ---
    add_section_header("Certifications & Technical Training")
    
    cert_data = [
        ("Google Cybersecurity Professional Certificate: ", "Covers Security Operations Center (SOC) workflows, SIEM packet analysis, threat detection, Linux/SQL security automation, and incident response frameworks."),
        ("Microsoft Networking and Cloud Computing: ", "Covers enterprise network architecture, TCP/IP fundamentals, cloud security models, subnetting, firewalls, and cloud infrastructure management."),
        ("Mastering Operating Systems: ", "In-depth training on process scheduling, memory virtualization, Unix/Linux system administration, privilege separation, and access control lists (ACLs)."),
        ("Cryptography & Network Security: ", "Hands-on implementation of symmetric/asymmetric encryption (AES, RSA), cryptographic hashing (SHA-256), public key infrastructure (PKI), and SSL/TLS security.")
    ]
    for c_title, c_desc in cert_data:
        p_c = doc.add_paragraph()
        add_bullet(p_c, c_title, c_desc)

    out_path = r"n:\Nidhish\project\Scanner\V2\Nidhish_V_Cybersecurity_Resume.docx"
    doc.save(out_path)
    print(f"Resume saved successfully to {out_path}")

if __name__ == "__main__":
    build_resume()
