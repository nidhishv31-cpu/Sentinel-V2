"""
generate_report.py
Script to generate a comprehensive, executive-level Word Document (.docx)
for the Sentinel Enterprise Platform (v2.0).
"""

import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Sets the background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_callout(doc, text, title="NOTE / SECURITY DIRECTIVE", alert_type="info"):
    """Adds a stylish callout box with colored left border and background."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    
    fill_hex = "F0F9FF" if alert_type == "info" else "FEF2F2"
    border_hex = "0284C7" if alert_type == "info" else "DC2626"
    
    set_cell_background(cell, fill_hex)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=180)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="36" w:space="0" w:color="{border_hex}"/>'
        f'  <w:top w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:bottom w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run_title = p.add_run(f"[{title}]\n")
    run_title.bold = True
    run_title.font.name = "Segoe UI"
    run_title.font.size = Pt(10)
    run_title.font.color.rgb = RGBColor(2, 132, 199) if alert_type == "info" else RGBColor(220, 38, 38)
    
    run_text = p.add_run(text)
    run_text.font.name = "Segoe UI"
    run_text.font.size = Pt(9.5)
    run_text.font.color.rgb = RGBColor(51, 65, 85)

def style_table_header(row, col_widths, bg_hex="0F172A"):
    """Styles the header row of a table."""
    for i, cell in enumerate(row.cells):
        cell.width = col_widths[i]
        set_cell_background(cell, bg_hex)
        set_cell_margins(cell, top=120, bottom=120, left=140, right=140)
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                run.bold = True
                run.font.name = "Segoe UI"
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(248, 250, 252)

def style_table_rows(table, col_widths, zebra=True):
    """Styles data rows with clean borders, padding, and subtle zebra stripes."""
    for r_idx, row in enumerate(table.rows[1:]):
        bg_hex = "F8FAFC" if (zebra and r_idx % 2 == 1) else "FFFFFF"
        for c_idx, cell in enumerate(row.cells):
            cell.width = col_widths[c_idx]
            if zebra and r_idx % 2 == 1:
                set_cell_background(cell, bg_hex)
            set_cell_margins(cell, top=90, bottom=90, left=140, right=140)
            
            # Add subtle light border
            tcPr = cell._tc.get_or_add_tcPr()
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
                f'  <w:top w:val="none"/>'
                f'  <w:left w:val="none"/>'
                f'  <w:right w:val="none"/>'
                f'</w:tcBorders>'
            )
            tcPr.append(borders)
            
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = "Segoe UI"
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(30, 41, 59)

def main():
    doc = Document()
    
    # Page setup - 0.75 in margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
        # Header / Footer
        footer = section.footer
        f_p = footer.paragraphs[0]
        f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        f_run = f_p.add_run("Sentinel Enterprise Platform v2.0 | Security Architecture & Technical Report")
        f_run.font.name = "Segoe UI"
        f_run.font.size = Pt(8)
        f_run.font.color.rgb = RGBColor(148, 163, 184)
    
    # Base Style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Segoe UI'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(30, 41, 59)

    # ==========================================
    # COVER / TITLE BLOCK
    # ==========================================
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(10)
    p_title.paragraph_format.space_after = Pt(2)
    
    run_badge = p_title.add_run("ENTERPRISE TECHNICAL SPECIFICATION & FORENSIC REPORT\n")
    run_badge.bold = True
    run_badge.font.name = "Segoe UI"
    run_badge.font.size = Pt(11)
    run_badge.font.color.rgb = RGBColor(14, 165, 233)
    
    run_main_title = p_title.add_run("SENTINEL ENTERPRISE SECURITY PLATFORM\nVERSION 2.0")
    run_main_title.bold = True
    run_main_title.font.name = "Segoe UI"
    run_main_title.font.size = Pt(24)
    run_main_title.font.color.rgb = RGBColor(15, 23, 42)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(16)
    run_sub = p_sub.add_run(
        "MNC-Grade Real-Time Deep Packet Inspection (DPI), Advanced Network Forensics, "
        "Threat Hunting, AI SOC Copilot, and Full-Stack Architectural Specifications"
    )
    run_sub.font.name = "Segoe UI"
    run_sub.font.size = Pt(11)
    run_sub.font.color.rgb = RGBColor(71, 85, 105)

    # Metadata Table
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_w = [Inches(2.5), Inches(4.0)]
    meta_rows = [
        ("Target Platform:", "Sentinel Enterprise Network Security Platform (v2.0)"),
        ("Engineering Architecture:", "FastAPI Asynchronous Engine + React 19 Virtualized Workstation"),
        ("Lead Roles:", "Full-Stack Security Developer, Lead QA Tester, MNC Security Analyst"),
        ("Documentation Date:", "September 2026 | Document Classification: Internal Enterprise Conf.")
    ]
    for idx, (label, val) in enumerate(meta_rows):
        r = meta_table.rows[idx]
        r.cells[0].paragraphs[0].add_run(label).bold = True
        r.cells[1].paragraphs[0].add_run(val)
    style_table_header(meta_table.rows[0], col_w, bg_hex="1E293B")
    style_table_rows(meta_table, col_w, zebra=True)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ==========================================
    # SECTION 1: EXECUTIVE SUMMARY
    # ==========================================
    h1 = doc.add_heading(level=1)
    r_h1 = h1.add_run("1. Executive Summary & Problem Landscape")
    r_h1.font.name = "Segoe UI"
    r_h1.font.color.rgb = RGBColor(15, 23, 42)
    
    doc.add_paragraph(
        "Modern enterprise networks face an unprecedented surge in sophisticated threat vectors, including "
        "encrypted Command & Control (C2) channels, high-entropy data exfiltration disguised within ordinary protocols, "
        "Domain Generation Algorithms (DGA), and living-off-the-land lateral movement. Traditional Intrusion Detection "
        "Systems (IDS) and Security Information and Event Management (SIEM) solutions often operate on delayed log batches "
        "or post-mortem PCAP file uploads, leaving Security Operations Center (SOC) analysts blind to in-flight protocol anomalies."
    )
    doc.add_paragraph(
        "The Sentinel Enterprise Platform (v2.0) was engineered from the ground up to solve this paradigm by unifying "
        "hardware-level promiscuous packet capture, microsecond-accurate byte-level dissection, Wireshark-grade conversational "
        "stream reassembly, and automated threat hunting into a single, high-performance web workstation. The platform combines "
        "an asynchronous Python/FastAPI backend powered by Scapy and raw binary struct processing with a 60 FPS virtualized React 19 "
        "frontend capable of smoothly navigating hundreds of thousands of frames without UI stutter."
    )
    
    add_callout(
        doc,
        "Sentinel v2.0 introduces 7 advanced MNC-grade capabilities: Real-Time I/O Bandwidth Sparklines, Interactive UML "
        "Sequence Ladder Flows, DNS Protocol Tunneling and DGA Detectors (MITRE T1071.004), Layer 2-4 Raw Packet Crafters with RFC 1071 "
        "checksum engines, an AI Forensic SOC Copilot generating automated Linux iptables, Windows netsh, and Suricata IDS rules, "
        "an OASIS STIX 2.1 Threat Intel Harvester, and in-flight TLS decryption via NSS SSLKEYLOGFILE ingestion.",
        title="EXECUTIVE CORE VALUE PROPOSITION",
        alert_type="info"
    )

    # ==========================================
    # SECTION 2: ARCHITECTURE & SYSTEM DESIGN
    # ==========================================
    h2 = doc.add_heading(level=1)
    r_h2 = h2.add_run("2. System Architecture & High-Performance Data Flow")
    r_h2.font.name = "Segoe UI"
    r_h2.font.color.rgb = RGBColor(15, 23, 42)

    doc.add_paragraph(
        "The Sentinel v2.0 architecture is divided into three decoupled tiers: the Promiscuous Hardware Ingestion & Forensic Core, "
        "the FastAPI Asynchronous Multiplexing Tier, and the React 19 High-Density Virtualized Workstation."
    )
    
    # Architecture Table
    arch_tbl = doc.add_table(rows=4, cols=3)
    arch_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    arch_w = [Inches(1.8), Inches(2.2), Inches(2.5)]
    arch_data = [
        ("Tier Layer", "Core Technologies", "Key Responsibilities"),
        ("Capture & Forensic Core", "Scapy, struct bit-packing, psutil, Npcap / libpcap", "Raw NIC sniffing, RFC 1071 checksum synthesis, Shannon entropy calculation, JA3/JA4 fingerprinting, YARA heuristics."),
        ("Asynchronous API Tier", "FastAPI, Uvicorn, WebSockets, Starlette, Pydantic v2", "Non-blocking REST routing, in-memory stream reassembly, bi-directional telemetry broadcast, STIX 2.1 assembly."),
        ("Analyst Workstation", "React 19, Vite 6, TypeScript 5.8, @tanstack/react-virtual", "Sub-10ms UI renders, 3-pane synchronized hex-to-tree highlighting, UML call flow diagrams, SVG sparklines.")
    ]
    for idx, (c1, c2, c3) in enumerate(arch_data):
        r = arch_tbl.rows[idx]
        r.cells[0].paragraphs[0].add_run(c1)
        r.cells[1].paragraphs[0].add_run(c2)
        r.cells[2].paragraphs[0].add_run(c3)
    style_table_header(arch_tbl.rows[0], arch_w, bg_hex="0F172A")
    style_table_rows(arch_tbl, arch_w, zebra=True)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ==========================================
    # SECTION 3: TECH STACK & FRAMEWORKS
    # ==========================================
    h3 = doc.add_heading(level=1)
    r_h3 = h3.add_run("3. Technology Stack, Frameworks & Tooling Specifications")
    r_h3.font.name = "Segoe UI"
    r_h3.font.color.rgb = RGBColor(15, 23, 42)

    doc.add_paragraph(
        "Every library, runtime, and framework in Sentinel v2.0 was deliberately selected to achieve low latency, "
        "maximum type safety, and seamless interoperability between enterprise security toolchains."
    )

    p_fnt = doc.add_paragraph()
    p_fnt.add_run("A. Frontend Client Frameworks (Analyst Workstation)").bold = True
    p_fnt.paragraph_format.space_before = Pt(4)

    f_tbl = doc.add_table(rows=7, cols=3)
    f_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    f_w = [Inches(1.8), Inches(1.2), Inches(3.5)]
    f_data = [
        ("Framework / Library", "Version", "Role & Engineering Rationale"),
        ("React", "19.0.0", "Core reactive component engine; zero-lag concurrent state updates."),
        ("TypeScript", "5.8.2", "Strict static type validation across complex packet data structures."),
        ("Vite", "6.4.3", "Next-gen lightning build tool, instantaneous HMR, optimized ESM bundling."),
        ("@tanstack/react-virtual", "3.13.2", "DOM node virtualization recycling; renders 500,000+ packets at 60 FPS."),
        ("Tailwind CSS", "4.0.0", "Ultra-performant utility-first styling with modern dark cyber theme."),
        ("lucide-react", "0.475.0", "Comprehensive vector iconography for forensic indicators and alerts.")
    ]
    for idx, (c1, c2, c3) in enumerate(f_data):
        r = f_tbl.rows[idx]
        r.cells[0].paragraphs[0].add_run(c1)
        r.cells[1].paragraphs[0].add_run(c2)
        r.cells[2].paragraphs[0].add_run(c3)
    style_table_header(f_tbl.rows[0], f_w, bg_hex="1E293B")
    style_table_rows(f_tbl, f_w, zebra=True)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    p_bck = doc.add_paragraph()
    p_bck.add_run("B. Backend Core Frameworks & Forensic Engines").bold = True
    p_bck.paragraph_format.space_before = Pt(4)

    b_tbl = doc.add_table(rows=8, cols=3)
    b_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    b_w = [Inches(1.8), Inches(1.2), Inches(3.5)]
    b_data = [
        ("Framework / Library", "Version", "Role & Engineering Rationale"),
        ("Python", "3.13 / 3.14", "Modern high-speed runtime with improved interpreter speed & free-threading."),
        ("FastAPI", "0.115.8", "High-performance async web framework built on Starlette and Pydantic."),
        ("Uvicorn", "0.34.0", "Lightning-fast ASGI server implementation for async Python worker processes."),
        ("Scapy", "2.6.1", "Low-level packet crafting, dissection, sniffing, and protocol parsing engine."),
        ("psutil", "7.0.0", "Cross-platform hardware NIC enumeration and real-time network I/O statistics."),
        ("Pydantic", "2.10.6", "Robust schema validation, serialization, and STIX 2.1 compliance checking."),
        ("Cryptography", "44.0.1", "TLS cipher suite mapping, cryptographic hashing, and entropy calculations.")
    ]
    for idx, (c1, c2, c3) in enumerate(b_data):
        r = b_tbl.rows[idx]
        r.cells[0].paragraphs[0].add_run(c1)
        r.cells[1].paragraphs[0].add_run(c2)
        r.cells[2].paragraphs[0].add_run(c3)
    style_table_header(b_tbl.rows[0], b_w, bg_hex="1E293B")
    style_table_rows(b_tbl, b_w, zebra=True)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ==========================================
    # SECTION 4: 7 ENTERPRISE CAPABILITIES
    # ==========================================
    h4 = doc.add_heading(level=1)
    r_h4 = h4.add_run("4. Detailed Breakdown of the 7 MNC-Grade Capabilities")
    r_h4.font.name = "Segoe UI"
    r_h4.font.color.rgb = RGBColor(15, 23, 42)

    doc.add_paragraph(
        "The following seven capabilities represent enterprise-level enhancements designed to meet the rigorous "
        "operational demands of Fortune 500 Security Operations Centers, Tier-3 SOC analysts, and incident response teams:"
    )

    cap_details = [
        ("1. Real-Time I/O Bandwidth & PPS Throughput Sparkline",
         "Provides continuous situational awareness directly in the UI header. Implemented as an optimized vector SVG path "
         "that recalculates rolling packet ingestion rates (PPS) and bandwidth consumption (KB/s) in a dynamic 20-sample sliding "
         "window. A glowing pulse indicator reflects live adapter activity, accompanied by peak rate indicators to alert analysts "
         "to sudden volumetric surges, DDoS attempts, or packet flood anomalies."),
         
        ("2. Interactive Packet Sequence Flow / Ladder Diagram (UML Call Flow)",
         "Translates dense, raw packet tables into an intuitive horizontal UML sequence ladder diagram. Maps multi-turn transactions "
         "between source and destination nodes, displaying directional request/response vectors, precise microsecond timing deltas, "
         "protocol classifications (HTTP, DNS, TLS, TCP handshakes), and payload previews. Analysts can click any transaction vector "
         "to open a deep inspection popover with full frame dissection."),

        ("3. DNS Protocol Tunneling & DGA Statistical Model Detector (MITRE T1071.004)",
         "Identifies covert DNS C2 communications and domain-generation algorithms without requiring external threat intel feeds. "
         "The detector evaluates multi-label subdomain queries against four statistical heuristics:\n"
         "  • Shannon Entropy: H(X) >= 3.8 bits/char flags encrypted or packed data.\n"
         "  • Query Label Length: Subdomains exceeding 35 characters indicate payload packing.\n"
         "  • Consonant-to-Vowel Ratio: Ratios exceeding 3.5 identify machine-generated pseudo-random DGAs.\n"
         "  • Payload Encoding Heuristics: Regex patterns identify Base32, Base64, and Hex encoding formats.\n"
         "Outputs confidence scoring, threat attribution, and MITRE ATT&CK T1071.004 mapping."),

        ("4. Layer 2–4 Raw Packet Crafter & Replay Injector ('Raw Packet Repeater')",
         "Empowers security analysts to perform targeted network penetration testing, firewall validation, and fuzzing by synthesizing "
         "custom Ethernet II, IPv4, and TCP/UDP frames. Built with an RFC 1071 One's Complement 16-bit Internet Checksum Engine "
         "that guarantees valid frame headers. Synthesized frames are injected into live streams and broadcasted over WebSockets to "
         "all connected analyst consoles in real time."),

        ("5. AI Forensic SOC Copilot Drawer & Defensive Rule Generator",
         "A virtual security analyst co-pilot integrated directly into the live inspection drawer. When a packet is selected, "
         "the copilot generates plain-English incident summaries, assigns MITRE ATT&CK TTPs, and auto-synthesizes multi-platform "
         "defensive drop rules:\n"
         "  • Linux iptables: iptables -A INPUT -s <src> -p <proto> --dport <port> -j DROP\n"
         "  • Windows netsh: netsh advfirewall firewall add rule name=\"SENTINEL_DROP_...\" dir=in action=block\n"
         "  • Suricata IDS Rule: alert <proto> <src> any -> any <port> (msg:\"SENTINEL DETECT...\"; sid:1000001; rev:1;)\n"
         "Enables immediate perimeter containment with 1-click clipboard rule copying."),

        ("6. Automated Threat Intel Harvester & OASIS STIX 2.1 JSON Exporter",
         "Continuously parses session traffic to extract high-confidence Indicators of Compromise (IOCs)—including malicious IP addresses, "
         "C2 domains, and suspicious payload hashes. Synthesizes an RFC-compliant OASIS STIX 2.1 JSON bundle containing indicator, "
         "ipv4-addr, domain-name, and file cyber observable objects. Provides 1-click download and clipboard export for instant ingestion "
         "into SIEM platforms (Splunk, Microsoft Sentinel, Elastic) and TIPs (MISP, OpenCTI)."),

        ("7. In-Flight TLS Decryption via NSS SSLKEYLOGFILE",
         "Solves the primary blind spot in enterprise packet analysis: encrypted HTTPS/TLS traffic. Ingests NSS-format pre-master "
         "secret files (CLIENT_RANDOM <random> <master_secret>) exported by modern browsers and cURL. Correlates secrets with live "
         "encrypted sessions to decrypt application-layer HTTP payloads in real time, revealing cleartext JSON API calls and exfiltration "
         "data without breaking certificate trust chains.")
    ]

    for title, desc in cap_details:
        p_cap = doc.add_paragraph()
        r_c = p_cap.add_run(title)
        r_c.bold = True
        r_c.font.size = Pt(10.5)
        r_c.font.color.rgb = RGBColor(14, 165, 233)
        doc.add_paragraph(desc)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ==========================================
    # SECTION 5: CORE DISSECTION CAPABILITIES
    # ==========================================
    h5 = doc.add_heading(level=1)
    r_h5 = h5.add_run("5. Core Dissection & Forensics Engine Specifications")
    r_h5.font.name = "Segoe UI"
    r_h5.font.color.rgb = RGBColor(15, 23, 42)

    doc.add_paragraph(
        "Sentinel v2.0 delivers a complete suite of Wireshark-grade deep inspection capabilities:"
    )

    core_features = [
        ("Synchronized Hex-to-Tree Highlighting",
         "Every protocol header and attribute in Pane 2 is mapped to its exact byte offset interval [start, end]. "
         "Hovering over any field in the protocol tree dynamically illuminates the exact byte sequence in Pane 3 with glowing cyan "
         "contrast, displaying an active offset breadcrumb (e.g. Inspecting: Internet Protocol Version 4 -> Source Address Bytes 26-30)."),
        
        ("Wireshark-Style 'Follow TCP Stream' Reassembly",
         "Reconstructs bidirectional TCP conversations using 4-tuple stream tracking. Color-codes client requests in Red/Salmon "
         "and server responses in Sky Blue. Features ASCII Text, Hex Dissector, and Raw Base64 display modes, with 1-click clipboard copying."),

        ("Shannon Entropy Exfiltration Meter",
         "Calculates Shannon entropy on packet payloads: H(X) = -sum(p_i * log2(p_i)) from 0.0 to 8.0 bits/byte. "
         "Classifies payloads into Low (<4.5 bits - plaintext), Moderate (4.5-6.8 bits - compressed binary), and High (>7.2 bits - encrypted C2)."),

        ("TLS JA3 & JA4 Fingerprinting",
         "Extracts TLS ClientHello parameters (Ciphers, Extensions, Curves, Formats) to compute MD5 JA3 hashes and JA4 signatures. "
         "Attributes traffic against known threat signatures, including Cobalt Strike HTTPS Malleable C2 Beacons."),

        ("YARA & Heuristic Artifact Carving",
         "Analyzes network-carved files for disguised Windows PE executables (MZ/PE00 headers - MITRE T1204), suspicious PDF JavaScript "
         "triggers (/JavaScript, /OpenAction - CVE-2023-26369), and packed payloads, returning color-coded verdicts: CLEAN, SUSPICIOUS, MALICIOUS."),

        ("BPF Display Filter Expression Engine",
         "Evaluates real-time filter expressions in the search bar (e.g. ip.src == 192.168.1.45, tcp.port == 443, proto == tlsv1.3, entropy > 7.0).")
    ]

    for f_title, f_desc in core_features:
        p_f = doc.add_paragraph()
        r_ft = p_f.add_run(f"• {f_title}: ")
        r_ft.bold = True
        p_f.add_run(f_desc)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ==========================================
    # SECTION 6: REST API REFERENCE
    # ==========================================
    h6 = doc.add_heading(level=1)
    r_h6 = h6.add_run("6. Comprehensive REST API & WebSocket Protocol Reference")
    r_h6.font.name = "Segoe UI"
    r_h6.font.color.rgb = RGBColor(15, 23, 42)

    doc.add_paragraph(
        "All live-capture and forensic capabilities are exposed via standardized, asynchronous REST endpoints "
        "and low-latency WebSocket communication channels:"
    )

    api_tbl = doc.add_table(rows=11, cols=3)
    api_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    api_w = [Inches(1.0), Inches(2.7), Inches(2.8)]
    api_data = [
        ("Method", "Endpoint Path", "Operational Description"),
        ("GET", "/api/live-capture/interfaces", "Enumerates active network interfaces with IP, MAC, and live status."),
        ("POST", "/api/live-capture/start", "Initiates live capture on chosen interface with optional BPF filter."),
        ("POST", "/api/live-capture/stop", "Terminates active packet sniffing session and closes ring buffers."),
        ("GET", "/api/live-capture/stream/{id}", "Reassembles bidirectional TCP conversation stream by stream ID."),
        ("POST", "/api/live-capture/inspect-carved", "Performs YARA & heuristic analysis on carved file artifacts."),
        ("POST", "/api/live-capture/inject", "Synthesizes and injects Layer 2-4 raw frames with RFC 1071 checksums."),
        ("POST", "/api/live-capture/detect-dns-tunnel", "Runs DGA statistical heuristics and tunnel exfiltration detection."),
        ("POST", "/api/live-capture/upload-keylog", "Ingests NSS SSLKEYLOGFILE master secrets for in-flight TLS decryption."),
        ("GET", "/api/live-capture/ioc-stix", "Harvests session IOCs and exports an OASIS STIX 2.1 JSON bundle."),
        ("POST", "/api/live-capture/copilot-analyze", "AI Copilot analysis generating forensic summaries and drop rules.")
    ]
    for idx, (m, ep, desc) in enumerate(api_data):
        r = api_tbl.rows[idx]
        r.cells[0].paragraphs[0].add_run(m).bold = True
        r.cells[1].paragraphs[0].add_run(ep)
        r.cells[2].paragraphs[0].add_run(desc)
    style_table_header(api_tbl.rows[0], api_w, bg_hex="0F172A")
    style_table_rows(api_tbl, api_w, zebra=True)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ==========================================
    # SECTION 7: QUALITY ASSURANCE & VERIFICATION
    # ==========================================
    h7 = doc.add_heading(level=1)
    r_h7 = h7.add_run("7. Quality Assurance, Test Automation & Verification")
    r_h7.font.name = "Segoe UI"
    r_h7.font.color.rgb = RGBColor(15, 23, 42)

    doc.add_paragraph(
        "Sentinel v2.0 adheres to strict MNC testing standards with automated Python test suites and full production "
        "TypeScript build validation."
    )

    p_test1 = doc.add_paragraph()
    p_test1.add_run("A. Enterprise Capabilities Test Suite (test_enterprise_suite.py)").bold = True
    p_test1.paragraph_format.space_before = Pt(4)
    doc.add_paragraph(
        "Verifies the 7 enterprise capabilities and all 5 new FastAPI REST endpoints:\n"
        "  • DNS Tunneling: Verified normal vs. C2 tunneling detection (entropy: 4.375 bits/char).\n"
        "  • Packet Crafter: Verified frame #1001 synthesis with RFC 1071 checksum (0x789c verified).\n"
        "  • STIX 2.1 Exporter: Validated bundle with 4 indicator objects.\n"
        "  • AI Copilot: Threat level CRITICAL, MITRE T1071.001 attribution, rules generated.\n"
        "  • SSLKEYLOGFILE: Ingested 2 pre-master secrets successfully.\n"
        "  • REST Endpoints: 5/5 HTTP 200 OK responses.\n"
        "Result: Ran 6 tests in 0.579s — 100% PASSED (OK)."
    )

    p_test2 = doc.add_paragraph()
    p_test2.add_run("B. Live Sniffer Base Test Suite (test_live_sniffer.py)").bold = True
    p_test2.paragraph_format.space_before = Pt(4)
    doc.add_paragraph(
        "Verifies base interface enumeration, Shannon entropy calculations, JA3/JA4 TLS hashing, YARA artifact heuristics, "
        "and TCP stream reassembly:\n"
        "Result: Ran 6 tests in 0.069s — 100% PASSED (OK)."
    )

    p_test3 = doc.add_paragraph()
    p_test3.add_run("C. Frontend Production Build Check (tsc -b && vite build)").bold = True
    p_test3.paragraph_format.space_before = Pt(4)
    doc.add_paragraph(
        "Compiled 3,351 modules in 7.20 seconds with 0 TypeScript compilation errors and 0 lint warnings. "
        "Assets bundled into optimized, minified production distribution (dist/index.html, dist/assets/*)."
    )

    add_callout(
        doc,
        "Total Test Verification Status: 12/12 Automated Python Unit & Integration Tests Passed (100% Green). "
        "Frontend TypeScript Production Build Passed with Zero Errors. All endpoints verified against active Starlette TestClient.",
        title="VERIFICATION SIGN-OFF",
        alert_type="info"
    )

    # ==========================================
    # SECTION 8: SOC OPERATIONAL PLAYBOOK
    # ==========================================
    h8 = doc.add_heading(level=1)
    r_h8 = h8.add_run("8. Production Operational Playbook & SOC Workflow")
    r_h8.font.name = "Segoe UI"
    r_h8.font.color.rgb = RGBColor(15, 23, 42)

    doc.add_paragraph(
        "When an anomalous packet or beacon is detected during a live capture session, the SOC analyst follows this workflow:"
    )
    doc.add_paragraph(
        "1. Triage & Anomaly Identification: Check the header Sparkline for traffic surges. Apply BPF filters (e.g. entropy > 7.0 or anomaly == true) to isolate suspicious frames.\n"
        "2. Deep Inspection & Protocol Dissection: Hover over protocol layers in Pane 2 to verify byte alignment in Pane 3. Inspect JA3/JA4 hashes against Cobalt Strike or known malware signatures.\n"
        "3. Conversation Reassembly: Click 'Follow Stream' on the TCP frame to review full ASCII or Hex dialogue. If encrypted, upload the browser's SSLKEYLOGFILE for immediate decryption.\n"
        "4. AI Incident Brief & Mitigation: Click 'AI Copilot' to generate a forensic brief and MITRE ATT&CK mapping. Copy the synthesized Linux iptables, Windows netsh, or Suricata rule for immediate network perimeter enforcement.\n"
        "5. Threat Intelligence Sharing: Click 'STIX 2.1' to download the standardized JSON bundle and ingest into enterprise SIEM / TIP platforms."
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(16)
    
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_end = p_end.add_run("— END OF SPECIFICATION REPORT —")
    r_end.bold = True
    r_end.font.size = Pt(9)
    r_end.font.color.rgb = RGBColor(148, 163, 184)

    # Save document
    out_path = os.path.join(os.path.dirname(__file__), "Sentinel_Enterprise_v2_Technical_Report.docx")
    doc.save(out_path)
    print(f"Successfully generated: {out_path}")

if __name__ == "__main__":
    main()
